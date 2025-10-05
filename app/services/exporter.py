from __future__ import annotations

import io
from pathlib import Path
from typing import Tuple

from docx import Document
from reportlab.lib.pagesizes import LETTER
from reportlab.pdfgen import canvas


def export_docx(brd_markdown: str) -> bytes:
    document = Document()
    for line in brd_markdown.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("```"):
            document.add_paragraph(line)
        else:
            document.add_paragraph(line)
    with io.BytesIO() as buffer:
        document.save(buffer)
        return buffer.getvalue()


def export_pdf(brd_markdown: str) -> bytes:
    buffer = io.BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=LETTER)
    width, height = LETTER
    x_margin, y_margin = 40, 40
    y_position = height - y_margin
    for line in brd_markdown.splitlines():
        if y_position <= y_margin:
            pdf.showPage()
            y_position = height - y_margin
        pdf.drawString(x_margin, y_position, line[:110])
        y_position -= 14
    pdf.save()
    return buffer.getvalue()


def export_brd(brd_markdown: str, export_format: str) -> Tuple[bytes, str]:
    if export_format == "docx":
        return export_docx(brd_markdown), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if export_format == "pdf":
        return export_pdf(brd_markdown), "application/pdf"
    raise ValueError("Unsupported export format")
